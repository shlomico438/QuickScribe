import os
import sys
import tempfile
import requests
from docx import Document
import smtplib
from email.message import EmailMessage
import torch
import omegaconf

# ==============================================================================
# 🔴 FIX: Whitelist ALL OmegaConf components blocking the load
# ==============================================================================
# 1. Force weights_only=False (Aggressive Monkey Patch)
original_load = torch.load
def unsafe_load(*args, **kwargs):
    # Forcefully overwrite the safety switch to False
    kwargs['weights_only'] = False 
    return original_load(*args, **kwargs)
torch.load = unsafe_load

# 2. Whitelist the specific classes mentioned in your error errors
torch.serialization.add_safe_globals([
    omegaconf.listconfig.ListConfig,      # Blocked in error 1
    omegaconf.dictconfig.DictConfig,      # Often blocked
    omegaconf.base.ContainerMetadata      # <--- THIS IS THE NEW ONE BLOCKING YOU
])
# ==============================================================================
import whisperx
import warnings
warnings.filterwarnings("ignore", category=UserWarning)

device = "cuda" if torch.cuda.is_available() else "cpu"
compute_type = "float16" if device == "cuda" else "float32"

# ---------- CONFIG ----------
SMTP_HOST = os.environ.get("SMTP_HOST")
SMTP_PORT = int(os.environ.get("SMTP_PORT", 587))
SMTP_USER = os.environ.get("SMTP_USER")
SMTP_PASS = os.environ.get("SMTP_PASS")
FROM_EMAIL = os.environ.get("FROM_EMAIL")

# ---------- INPUT HANDLING ----------
def get_inputs():
    # Environment variables for serverless mode
    customer_email = os.environ.get("CUSTOMER_EMAIL")
    zoom_url = os.environ.get("ZOOM_URL")
    local_file = os.environ.get("LOCAL_FILE")  # optional local file

    if customer_email and (zoom_url or local_file):
        return customer_email, zoom_url, local_file

    # Interactive mode
    print("Interactive mode: provide input manually.")
    choice = input("Do you want to use a local file or Zoom URL? [local/zoom]: ").strip().lower()
    if choice == "local":
        local_file = input("Enter local MP3/MP4 path: ").strip()
        zoom_url = None
    else:
        zoom_url = input("Enter Zoom recording URL: ").strip()
        local_file = None
    customer_email = input("Enter customer email (leave blank to skip sending email): ").strip()
    if customer_email == "":
        customer_email = None
    return customer_email, zoom_url, local_file

if device == "cuda":
    CUSTOMER_EMAIL, ZOOM_URL, LOCAL_FILE = get_inputs()
else:
    # --- DEBUGGING DEFAULTS ---
    # We will map your Windows folder to this Linux folder
    LOCAL_FILE = "/data/small_test.mp3" 
    CUSTOMER_EMAIL = None#"shlomico1234@gmail.com"
    ZOOM_URL = None

if not (ZOOM_URL or LOCAL_FILE):
    print({"status": "error", "message": "No input file provided"})
    sys.exit(1)

# ---------- STEP 1: Get audio ----------
try:
    if LOCAL_FILE:
        temp_audio_file_path = LOCAL_FILE
    else:
        response = requests.get(ZOOM_URL)
        response.raise_for_status()
        temp_audio_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
        temp_audio_file.write(response.content)
        temp_audio_file.close()
        temp_audio_file_path = temp_audio_file.name
except Exception as e:
    print({"status": "error", "message": f"Failed to get audio: {e}"})
    sys.exit(1)

# ---------- STEP 2: Transcribe ----------
try:
    # Load full model only on GPU, smallest model on CPU
    if device == "cuda":
        print("GPU detected: loading full WhisperX large-v2 model...")
        model = whisperx.load_model("large-v2", device=device, compute_type=compute_type)
    else:
        print("No GPU detected: loading smallest WhisperX model for CPU debugging...")
        model = whisperx.load_model("tiny", device=device, compute_type=compute_type)
        
    audio = whisperx.load_audio(temp_audio_file_path)
    result = model.transcribe(audio, batch_size=16 if device == "cuda" else 4) # verbose removed to clean output
    transcript_text = ""
    
    # Improved text extraction loop
    for segment in result['segments']:
        transcript_text += segment['text'] + "\n"

except Exception as e:
    print({"status": "error", "message": f"Transcription failed: {e}"})
    sys.exit(1)

# ---------- STEP 3: Generate DOCX ----------
try:
    doc = Document()
    doc.add_heading("Transcript", 0)
    doc.add_paragraph(transcript_text)
    
    # Fix output path for both Windows (Local) and Linux (RunPod)
    if os.name == 'nt': # Windows
        output_dir = r"C:\Work\runpod\QuickScribe\Output"
    else: # Linux/RunPod
        output_dir = "/input/output"
        
    os.makedirs(output_dir, exist_ok=True)
    docx_file_path = os.path.join(output_dir, "transcript.docx")
    doc.save(docx_file_path)
except Exception as e:
    print({"status": "error", "message": f"DOCX generation failed: {e}"})
    sys.exit(1)

# ---------- STEP 4: Deliver output ----------
def send_email(file_path, to_email):
    try:
        msg = EmailMessage()
        msg['Subject'] = 'Your Transcript'
        msg['From'] = FROM_EMAIL
        msg['To'] = to_email

        with open(file_path, 'rb') as f:
            file_data = f.read()

        msg.add_attachment(
            file_data,
            maintype='application',
            subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename='transcript.docx'
        )

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as smtp:
            smtp.starttls()
            smtp.login(SMTP_USER, SMTP_PASS)
            smtp.send_message(msg)
        return True
    except Exception as e:
        print({"status": "error", "message": f"Email sending failed: {e}"})
        return False

if CUSTOMER_EMAIL:
    success = send_email(docx_file_path, CUSTOMER_EMAIL)
    if success:
        print({"status": "success", "message": f"Transcript sent to {CUSTOMER_EMAIL}"})
else:
    print(f"No email provided. Transcript saved at: {docx_file_path}")